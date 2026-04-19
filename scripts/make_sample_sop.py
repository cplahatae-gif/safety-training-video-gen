"""Generate a tiny sample SOP DOCX for end-to-end smoke testing."""
from pathlib import Path
from docx import Document

OUT = Path(__file__).parent.parent / "samples" / "고소작업차_아웃리거_점검.docx"

doc = Document()
doc.add_heading("고소작업차 아웃리거 점검 표준작업지침", level=1)

doc.add_heading("1. 적용 범위", level=2)
doc.add_paragraph("본 지침은 건설 현장에서 고소작업차(차량탑재형 고소작업대)를 운용하는 모든 작업자에게 적용한다.")

doc.add_heading("2. 법적 근거", level=2)
doc.add_paragraph("산업안전보건기준에 관한 규칙 제186조 (고소작업대 설치 등의 조치)")
doc.add_paragraph("산업안전보건기준에 관한 규칙 제189조 (작업대 사용 시 조치)")

doc.add_heading("3. 주요 위험 요인", level=2)
doc.add_paragraph("• 아웃리거 미전개 또는 부분 전개로 인한 차량 전도 (사망 위험)")
doc.add_paragraph("• 작업대 탑승 중 안전대 미착용으로 인한 추락 (사망 위험)")
doc.add_paragraph("• 연약 지반에서 받침판 미사용으로 인한 아웃리거 침하 및 전복")

doc.add_heading("4. 작업 전 점검 절차", level=2)
doc.add_paragraph("1단계: 작업 지반을 확인한다. 경사도 5도 이하, 단단한 노면에서만 설치한다.")
doc.add_paragraph("2단계: 아웃리거 4개 모두를 지면에 완전히 닿도록 전개하고 받침판을 반드시 사용한다.")
doc.add_paragraph("3단계: 수평계를 확인하여 차체 수평을 맞춘다. 타이어가 지면에서 떠야 한다.")
doc.add_paragraph("4단계: 작업자는 탑승 전 안전대를 작업대 내 고정 지점에 체결한다.")
doc.add_paragraph("5단계: 상승 전 주변 장애물(전선, 구조물)과의 이격거리를 확인한다.")

doc.add_heading("5. 작업 대상", level=2)
doc.add_paragraph("고소작업차 운전자 및 보조 작업자")

doc.add_heading("6. 자주 발생하는 위반 사항", level=2)
doc.add_paragraph("• 아웃리거를 완전히 전개하지 않고 작업 시작")
doc.add_paragraph("• 작업대 내에서 안전대를 체결하지 않음")
doc.add_paragraph("• 받침판 없이 연약 지반에 설치")
doc.add_paragraph("• 정격 하중 초과 탑승")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
